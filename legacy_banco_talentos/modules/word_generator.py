from docx import Document
from docx.shared import Pt, RGBColor
import json
import os
import tempfile
from datetime import datetime
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import unicodedata

def nome_arquivo_seguro(valor):
    valor = str(valor or "curriculo")
    valor = unicodedata.normalize("NFD", valor)
    valor = valor.encode("ascii", "ignore").decode("ascii")
    valor = re.sub(r"[^a-zA-Z0-9._-]+", "_", valor)
    valor = valor.strip("_")
    return valor[:120] or "curriculo"

class WordGenerator:
    
    def __init__(self, templates_dir="templates"):
        self.templates_dir = templates_dir
        self.config = self.load_config()


    @staticmethod
    def formatar_campo_multilinha(valor, usar_bullet=True):
        if not valor:
            return ""

        valor = str(valor).strip()

        if not valor:
            return ""

        if "\n" in valor:
            linhas = [linha.strip(" -•\t") for linha in valor.splitlines() if linha.strip()]
        else:
            linhas = [parte.strip(" -•\t") for parte in valor.split("|") if parte.strip()]

        if not linhas:
            return ""

        if usar_bullet:
            return "\n".join(f"• {linha}" for linha in linhas)

        return "\n".join(linhas)

    @staticmethod
    def add_hyperlink(paragraph, text, url):
        """
        Adiciona um hyperlink clicável em um parágrafo do Word.
        """

        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True
        )

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")

        r_pr = OxmlElement("w:rPr")

        # Estilo visual de link: azul e sublinhado
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        r_pr.append(color)

        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        r_pr.append(underline)

        new_run.append(r_pr)

        text_element = OxmlElement("w:t")
        text_element.text = text
        new_run.append(text_element)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

        return hyperlink

    @staticmethod
    def limpar_runs(paragraph):
        """
        Remove todos os runs de um parágrafo.
        """
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)

    @staticmethod
    def substituir_placeholder_por_hyperlink(paragraph, placeholder, url, texto_link=None):
        """
        Substitui um placeholder por hyperlink.
        Se não existir URL, remove a linha inteira do LinkedIn.
        """

        texto_paragrafo = paragraph.text

        if placeholder not in texto_paragrafo:
            return False

        url = str(url or "").strip()

        # Se não tiver LinkedIn, remove a linha inteira:
        # Exemplo: "➜ Linkedin: LINKEDIN"
        if not url:
            p = paragraph._element
            parent = p.getparent()

            if parent is not None:
                parent.remove(p)

            return True

        texto_antes, texto_depois = texto_paragrafo.split(placeholder, 1)

        WordGenerator.limpar_runs(paragraph)

        if texto_antes:
            paragraph.add_run(texto_antes)

        WordGenerator.add_hyperlink(
            paragraph,
            texto_link or url,
            url
        )

        if texto_depois:
            paragraph.add_run(texto_depois)

        return True

    @staticmethod
    def substituir_linkedin_no_documento(doc, linkedin_url):
        """
        Procura o placeholder LINKEDIN no documento e substitui por hyperlink.
        """

        texto_link = "Abrir perfil no LinkedIn"

        # Parágrafos fora de tabela
        for paragraph in doc.paragraphs:
            WordGenerator.substituir_placeholder_por_hyperlink(
                paragraph,
                "LINKEDIN",
                linkedin_url,
                texto_link
            )

        # Parágrafos dentro de tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        WordGenerator.substituir_placeholder_por_hyperlink(
                            paragraph,
                            "LINKEDIN",
                            linkedin_url,
                            texto_link
                        )

    def load_config(self):
        import os

        base_dir = os.path.dirname(os.path.abspath(__file__))
        config_path = os.path.join(base_dir, "..", "templates", "config.json")

        with open(config_path, "r", encoding="utf-8") as f:
            return json.load(f)
    
    def get_template_info(self, template_id):
        """Retorna informações do template"""
        for template in self.config['templates']:
            if template['id'] == template_id:
                return template
        return None
    
    def generate_document(self, template_id, candidate_data):
        """
        Gera documento Word preenchido com dados do candidato
        
        Args:
            template_id: ID do template (ex: 'cv', 'carta')
            candidate_data: Dicionário com dados do candidato
        
        Returns:
            Caminho do arquivo gerado
        """
        # Obter informações do template
        template_info = self.get_template_info(template_id)
        if not template_info:
            raise ValueError(f"Template '{template_id}' não encontrado")
        
        candidate_data = dict(candidate_data)

        campos_multilinha = [
            "formacao_academica",
            "cursos_certificacoes",
            "conhecimento_tecnico",
            "experiencia_profissional"
        ]

        for campo in campos_multilinha:
            candidate_data[campo] = WordGenerator.formatar_campo_multilinha(
                candidate_data.get(campo, ""),
                usar_bullet=True)


        print(f'Caminho arquivo:{self.templates_dir}')
        # Carregar template
        template_path = os.path.join(self.templates_dir,template_info['arquivo'])

        doc = Document(template_path)

        placeholders = template_info['placeholders']
        print(f'placeholders: {placeholders}')
        print(f'candidate_data: {candidate_data}')
        for placeholder, field_name in placeholders.items():
            if placeholder.upper() == "LINKEDIN":
                continue

            value = candidate_data.get(field_name, "")
            self._replace_text_in_document(doc, placeholder, str(value))

        linkedin_url = candidate_data.get("linkedin", "") or candidate_data.get("Link_Linkedin", "")
        WordGenerator.substituir_linkedin_no_documento(doc, linkedin_url)
        
        # Salvar documento
        nome_candidato = nome_arquivo_seguro(candidate_data.get("nome", "curriculo"))
        template_seguro = nome_arquivo_seguro(template_id)

        output_filename = f"{nome_candidato}_{template_seguro}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_dir = tempfile.gettempdir()
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, output_filename)
        doc.save(output_path)
        
        return output_path, output_filename
    
    def _replace_text_in_document(self, doc, placeholder, replacement):
        """Substitui placeholder no documento"""
        # Substituir em parágrafos
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                self._replace_text_in_paragraph(paragraph, placeholder, replacement)
        
        # Substituir em tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            self._replace_text_in_paragraph(paragraph, placeholder, replacement)
    
    def _replace_text_in_paragraph(self, paragraph, placeholder, replacement):
        full_text = paragraph.text

        if placeholder not in full_text:
            return

        if replacement is None or str(replacement).strip() == "":
            p = paragraph._element
            p.getparent().remove(p)
            return

        new_text = full_text.replace(placeholder, str(replacement))

        for run in paragraph.runs:
            run.text = ""

        if paragraph.runs:
            run = paragraph.runs[0]
        else:
            run = paragraph.add_run()

        linhas = new_text.split("\n")

        run.text = linhas[0]

        for linha in linhas[1:]:
            run.add_break()
            run.add_text(linha)
    def list_templates(self):
        """Retorna lista de templates disponíveis"""
        return [
            {
                "id": t['id'],
                "nome": t['nome']
            }
            for t in self.config['templates']
        ]