"""
Módulo para extração de dados de CVs usando Gemini API
"""
import re
import json
import time
import logging
from pathlib import Path
from google.genai import types
from docx import Document
import PyPDF2
from config.gemini import get_gemini
from utils.file_handler import FileHandler

logger = logging.getLogger(__name__)

class GeminiExtractor:
    """Extrai dados de CVs usando Gemini API"""
    
    PROMPT_TEMPLATE = """Extraia os dados de um currículo a partir do documento fornecido.

RETORNE APENAS UM JSON VÁLIDO, SEM TEXTO ADICIONAL, SEM EXPLICAÇÕES, SEM MARKDOWN.

O JSON DEVE CONTER EXATAMENTE AS SEGUINTES CHAVES (todas obrigatórias):
- Nome
- Nacionalidade
- Estado Civil
- Idade
- Endereco
- Telefone
- Email
- Link_Linkedin
- Skil
- Formacao_Academica
- Nivel_Idioma_Ingles
- Nivel_Idioma_Espanhol
- Cursos_Certificacoes
- Conhecimento_Tecnico
- Experiencia_Profissional

REGRAS IMPORTANTES:

1. O retorno deve ser um objeto JSON (dict), nunca lista.
2. Todas as chaves devem existir, mesmo que vazias.
3. Se não encontrar um valor, retornar "" (string vazia).
4. Cada chave deve conter apenas um valor do tipo string.
5. Não incluir listas, objetos aninhados ou múltiplos valores.
6. Não incluir quebras de linha no JSON.
7. Não incluir ```json ou ```.

CRITÉRIOS DE EXTRAÇÃO:

Nome:
- Nome completo do candidato
- Localizado no topo do currículo

Nacionalidade:
- Nacionalidade do Canditado, caso não tenha nada considerar Brasileiro

Estado_Civil:
- Estado Civil do Canditado, caso não tenha nada considerar vazio

Idade:
- Caso não exista, calcular a partir da data de nascimento

Endereco:
- Cidade, estado ou endereço completo

Telefone:
- Numero de Telefone no formato ddd-xxxxx-xxxx

Email:
- Email
- Caso não exista, deixar em branco

Link_Linkedin:
- Capturar somente o link do perfil LinkedIn do candidato.
- Considerar links nos formatos:
  - https://www.linkedin.com/in/nome-do-perfil
  - http://www.linkedin.com/in/nome-do-perfil
  - www.linkedin.com/in/nome-do-perfil
  - linkedin.com/in/nome-do-perfil
  - /in/nome-do-perfil
- Se encontrar apenas "/in/nome-do-perfil", completar como:
  https://www.linkedin.com/in/nome-do-perfil
- Se encontrar "linkedin.com/in/nome-do-perfil", completar como:
  https://www.linkedin.com/in/nome-do-perfil
- Retornar sempre uma URL completa iniciando com "https://".
- Não retornar texto adicional, emojis ou descrição.
- Se não encontrar um perfil LinkedIn válido, retornar "".

Skil:
- Resumo das qualificações profissionais

Formacao_Academica:
- Listar cada formação separada por " | "
- Formato: Curso - Instituição (Período ou Ano)
- Se houver apenas uma formação, retornar apenas ela
- Não usar listas JSON

Nivel_Idioma_Ingles:
- Nivel_Idioma_Ingles
- Caso não exista, deixar em branco

Nivel_Idioma_Espanhol:
- Nivel_Idioma_Espanhol
- Caso não exista, deixar em branco

Cursos_Certificacoes:
- Listar cada curso ou certificação separada por " | "
- Formato: Nome do curso/certificação - Instituição (Ano), quando existir
- Não usar listas JSON

Conhecimento_Tecnico:
- Listar os conhecimentos técnicos agrupados por tema, separados por " | "
- Exemplo: Linguagens: Python, SQL | BI: Power BI, Tableau | Cloud: AWS, Azure
- Não usar listas JSON

Experiencia_Profissional:
- Listar cada experiência separada por " | "
- Formato obrigatório:
  Empresa - Cargo (Período)
- Exemplo:
  Algar - Analista de Negócios TI e Telecom (2025) | Reply - Analista de Negócios TI (2020 a 2025) | Truckpad - Analista de Qualidade de Software (2019 a 2020)
- Não juntar empresas diferentes no mesmo bloco
- Não usar listas JSON

FORMATO FINAL OBRIGATÓRIO:
{{
  "Nome": "",
  "Nacionalidade": "",
  "Estado_Civil": "",
  "Idade": "",
  "Endereco": "",
  "Telefone": "",
  "Email": "",
  "Link_Linkedin": "",
  "Skil": "",
  "Formacao_Academica": "",
  "Nivel_Idioma_Ingles": "",
  "Nivel_Idioma_Espanhol": "",
  "Cursos_Certificacoes": "",
  "Conhecimento_Tecnico": "",
  "Experiencia_Profissional": ""
}}"""
    
    def __init__(self):
        """Inicializa o extrator Gemini"""
        self.gemini_config = None
        self.client = None
        self.model = None
        self.last_error = None
        try:
            self.gemini_config = get_gemini()
            self.client = self.gemini_config.get_client()
            self.model = self.gemini_config.get_model()
        except Exception as e:
            tipo_erro = self._classificar_exception(e)
            self._set_error(tipo_erro, f"Gemini indisponivel: {e}")
            logger.warning(
                f"Gemini indisponivel. Extracao local sera usada como fallback: {e}"
            )

    def _normalizar_campo_multilinha(self, valor, usar_bullet=True):
        """
        Converte textos separados por | em múltiplas linhas.
        Exemplo:
        Empresa A - Cargo | Empresa B - Cargo
        vira:
        • Empresa A - Cargo
        • Empresa B - Cargo
        """

        if not valor:
            return ""

        valor = str(valor).strip()

        if not valor:
            return ""

        # Se já vier com quebra de linha, apenas limpa espaços
        if "\n" in valor:
            linhas = [linha.strip(" -•\t") for linha in valor.splitlines() if linha.strip()]
        else:
            linhas = [parte.strip(" -•\t") for parte in re.split(r"\s*\|\s*", valor) if parte.strip()]

        if not linhas:
            return ""

        if usar_bullet:
            return "\n".join(f"• {linha}" for linha in linhas)

        return "\n".join(linhas)

    def _normalizar_linkedin(self, valor):
        """
        Normaliza o link do LinkedIn para URL completa.
        """

        if not valor:
            return ""

        valor = str(valor).strip()

        if not valor:
            return ""

        # Remove espaços internos comuns
        valor = valor.replace(" ", "")

        # Remove caracteres comuns no final
        valor = valor.rstrip(".,;|)")

        # Se vier apenas /in/perfil
        if valor.startswith("/in/"):
            return f"https://www.linkedin.com{valor}"

        # Se vier linkedin.com/in/perfil
        if valor.startswith("linkedin.com/in/"):
            return f"https://www.{valor}"

        # Se vier www.linkedin.com/in/perfil
        if valor.startswith("www.linkedin.com/in/"):
            return f"https://{valor}"

        # Se vier http://linkedin.com/in/perfil
        if valor.startswith("http://linkedin.com/in/"):
            valor = valor.replace("http://linkedin.com", "https://www.linkedin.com", 1)
            return valor

        # Se vier https://linkedin.com/in/perfil
        if valor.startswith("https://linkedin.com/in/"):
            valor = valor.replace("https://linkedin.com", "https://www.linkedin.com", 1)
            return valor

        # Se vier http://www.linkedin.com/in/perfil
        if valor.startswith("http://www.linkedin.com/in/"):
            valor = valor.replace("http://", "https://", 1)
            return valor

        # Se já estiver certo
        if valor.startswith("https://www.linkedin.com/in/"):
            return valor

        # Se não for perfil LinkedIn válido
        return ""
    def _extrair_retry_delay(self, exception, default_delay=60):
        """
        Extrai do erro Gemini o tempo sugerido em:
        'Please retry in 48.931894276s'
        """
        msg = str(exception)

        match = re.search(r"retry in ([\d.]+)s", msg, re.IGNORECASE)

        if match:
            return int(float(match.group(1))) + 2

        return default_delay

    def _set_error(self, tipo, mensagem):
        self.last_error = {
            "tipo": tipo,
            "mensagem": mensagem
        }

    def _classificar_exception(self, exception):
        msg = str(exception).lower()

        if "503" in msg or "unavailable" in msg or "high demand" in msg or "overloaded" in msg:
            return "gemini_modelo_indisponivel_alta_demanda"

        if "429" in msg or "quota" in msg or "rate limit" in msg:
            return "gemini_limite_quota"

        if "unsupported mime" in msg or "mime" in msg:
            return "gemini_mime_invalido"

        if "json" in msg or isinstance(exception, json.JSONDecodeError):
            return "gemini_json_invalido"

        if "timeout" in msg:
            return "gemini_timeout"

        if "permission" in msg or "unauthorized" in msg or "403" in msg or "401" in msg:
            return "gemini_erro_autenticacao"

        return "gemini_erro_api"

    def _normalize_json_response(self, dados):
        """Normaliza a resposta do modelo Gemini"""
        if isinstance(dados, list):
            if len(dados) == 1 and isinstance(dados[0], dict):
                return dados[0]
            else:
                return dados
        elif isinstance(dados, dict):
            return dados
        else:
            raise ValueError("Formato inesperado de dados retornados.")
    
    def extract_from_pdf(self, pdf_path, max_tentativas=10, delay_retry=10):
        """
        Extrai dados de um PDF usando Gemini
        
        Args:
            pdf_path (str): Caminho do arquivo PDF
            max_tentativas (int): Número máximo de tentativas
            delay_retry (int): Tempo entre tentativas em segundos
            
        Returns:
            dict: Dados extraídos ou None em caso de erro
        """
        
        if not self.client or not self.model:
            logger.warning("Gemini nao configurado. Usando fallback local.")
            return None

        for tentativa in range(1, max_tentativas + 1):
            try:
                logger.info(f" Processando PDF (tentativa {tentativa}/{max_tentativas}): {pdf_path}")
                
                # Fazer upload do PDF
                uploaded_file = self.client.files.upload(file=pdf_path)
                logger.info(f" PDF enviado para Gemini")
                
                # Chamar o modelo Gemini
                response = self.client.models.generate_content(
                    model=self.model,
                    contents=[uploaded_file, self.PROMPT_TEMPLATE],
                    config=types.GenerateContentConfig(
                        response_mime_type="application/json"
                    )
                )
                
                # Parsear a resposta
                dados_brutos = response.text.strip()
                
                # Limpar bordas de markdown se existirem
                dados_limpos = dados_brutos.strip('```json').strip('```').strip()
                
                # Converter para JSON
                json_dados = json.loads(dados_limpos)
                dados = self._normalize_json_response(json_dados)
                campos_multilinha = [
                    "Formacao_Academica",
                    "Cursos_Certificacoes",
                    "Conhecimento_Tecnico",
                    "Experiencia_Profissional"]

                for campo in campos_multilinha:
                    dados[campo] = self._normalizar_campo_multilinha(dados.get(campo, ""),usar_bullet=True)

                dados["Link_Linkedin"] = self._normalizar_linkedin(dados.get("Link_Linkedin", ""))
                
                logger.info(f" Dados extraídos com sucesso")
                return dados
                
            except Exception as e:
                tipo_erro = self._classificar_exception(e)
                self._set_error(
                    tipo_erro,
                    f"Erro na tentativa {tentativa}/{max_tentativas} ao extrair PDF: {e}"
                )

                logger.error(f" Erro na tentativa {tentativa}: {e}")

                if tipo_erro in {"gemini_limite_quota", "gemini_erro_autenticacao"}:
                    logger.warning(
                        "Falha nao recuperavel nesta execucao. "
                        "Usando extracao local sem aguardar novas tentativas."
                    )
                    return None
                
                if tentativa < max_tentativas:
                    if tipo_erro == "gemini_limite_quota":
                        tempo_espera = self._extrair_retry_delay(e, default_delay=60)
                    elif tipo_erro == "gemini_modelo_indisponivel_alta_demanda":
                        tempo_espera = delay_retry * tentativa
                    else:
                        tempo_espera = delay_retry * tentativa

                    logger.info(
                        f" Tentando novamente em {tempo_espera}s... "
                        f"Motivo: {tipo_erro}"
                    )

                    time.sleep(tempo_espera)
                else:
                    logger.error(f" Limite de tentativas atingido para {pdf_path}")
                    return None

    def _extract_text_from_pdf(self, pdf_path):
        text_parts = []
        with open(pdf_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
        return "\n".join(text_parts)

    def _extract_text_from_docx(self, docx_path):
        document = Document(docx_path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts)

    def _extract_text_local(self, file_path):
        suffix = Path(file_path).suffix.lower()
        if suffix == ".pdf":
            return self._extract_text_from_pdf(file_path)
        if suffix == ".docx":
            return self._extract_text_from_docx(file_path)

        pdf_path = FileHandler.ensure_pdf(file_path)
        if pdf_path:
            try:
                return self._extract_text_from_pdf(pdf_path)
            finally:
                if pdf_path != file_path:
                    FileHandler.delete_file(pdf_path)

        return ""

    def _normalize_for_match(self, value):
        import unicodedata

        normalized = unicodedata.normalize("NFD", str(value or ""))
        normalized = "".join(char for char in normalized if unicodedata.category(char) != "Mn")
        return re.sub(r"\s+", " ", normalized).strip().lower()

    def _clean_heading_key(self, value):
        normalized = self._normalize_for_match(value)
        normalized = re.sub(r"[^a-z0-9]+", " ", normalized)
        return re.sub(r"\s+", " ", normalized).strip()

    def _is_heading_match(self, line, aliases):
        if not line or len(line) > 120:
            return False

        key = self._clean_heading_key(line)
        alias_keys = {self._clean_heading_key(alias) for alias in aliases}

        if key in alias_keys:
            return True

        normalized = self._normalize_for_match(line)
        for alias in aliases:
            alias_norm = self._normalize_for_match(alias)
            if re.match(rf"^\s*{re.escape(alias_norm)}\s*[:\-|]", normalized):
                return True

        return False

    def _strip_heading_prefix(self, line, aliases):
        normalized_line = self._normalize_for_match(line)
        for alias in sorted(aliases, key=len, reverse=True):
            alias_norm = self._normalize_for_match(alias)
            match = re.match(rf"^\s*{re.escape(alias_norm)}\s*[:\-|]\s*(.+)$", normalized_line)
            if match:
                delimiter = re.search(r"[:\-|]", line)
                if delimiter:
                    return line[delimiter.end():].strip()
        return ""

    def _compact_multiline(self, value, max_chars=6000):
        lines = [
            re.sub(r"\s+", " ", line).strip(" -\t")
            for line in str(value or "").splitlines()
            if re.sub(r"\s+", " ", line).strip(" -\t")
        ]
        text = "\n".join(lines)
        return text[:max_chars].strip()

    def _extract_section(self, clean_lines, target_aliases, all_heading_aliases, max_chars=6000):
        parts = []
        collecting = False

        for line in clean_lines:
            target_hit = self._is_heading_match(line, target_aliases)
            any_heading_hit = self._is_heading_match(line, all_heading_aliases)

            if target_hit:
                collecting = True
                inline_value = self._strip_heading_prefix(line, target_aliases)
                if inline_value:
                    parts.append(inline_value)
                continue

            if collecting and any_heading_hit:
                break

            if collecting:
                parts.append(line)

        return self._compact_multiline("\n".join(parts), max_chars=max_chars)

    def _join_unique_sections(self, sections, max_chars=7000):
        seen = set()
        result = []
        for section in sections:
            for line in str(section or "").splitlines():
                clean = re.sub(r"\s+", " ", line).strip()
                key = self._clean_heading_key(clean)
                if clean and key not in seen:
                    seen.add(key)
                    result.append(clean)
        return "\n".join(result)[:max_chars].strip()

    def _extract_keyword_lines(self, clean_lines, keywords, max_lines=40, max_chars=6000):
        keyword_keys = [self._normalize_for_match(keyword) for keyword in keywords]
        matches = []
        for line in clean_lines:
            line_norm = self._normalize_for_match(line)
            if any(keyword in line_norm for keyword in keyword_keys):
                matches.append(line)
            if len(matches) >= max_lines:
                break
        return self._join_unique_sections(matches, max_chars=max_chars)

    def _extract_label_value(self, clean_lines, labels):
        label_keys = [self._normalize_for_match(label) for label in labels]
        for line in clean_lines:
            line_norm = self._normalize_for_match(line)
            for label in label_keys:
                if not line_norm.startswith(label):
                    continue
                match = re.search(r"[:\-|]\s*(.+)$", line)
                if match:
                    return match.group(1).strip()
                return line.strip()
        return ""

    def _extract_language_level(self, clean_lines, labels):
        label_keys = [self._normalize_for_match(label) for label in labels]
        for line in clean_lines:
            line_norm = self._normalize_for_match(line)
            if any(label in line_norm for label in label_keys):
                match = re.search(r"[:\-|]\s*(.+)$", line)
                return (match.group(1) if match else line).strip()
        return ""

    def _fallback_extract_from_text(self, text, file_path):
        raw_text = str(text or "")
        clean_lines = [
            re.sub(r"\s+", " ", line).strip()
            for line in raw_text.splitlines()
            if re.sub(r"\s+", " ", line).strip()
        ]
        compact_text = re.sub(r"\s+", " ", raw_text).strip()

        email_match = re.search(r"[\w.+-]+@[\w-]+(?:\.[\w-]+)+", compact_text)
        phone_match = re.search(
            r"(?:\+?55\s*)?(?:\(?\d{2}\)?\s*)?(?:9\s*)?\d{4}[-.\s]?\d{4}",
            compact_text,
        )
        linkedin_match = re.search(
            r"(?:https?://)?(?:www\.)?linkedin\.com/in/[A-Za-z0-9%_\-./]+|/in/[A-Za-z0-9%_\-./]+",
            compact_text,
            re.IGNORECASE,
        )

        skip_name_patterns = re.compile(
            r"curr[ií]culo|resume|linkedin|e-mail|email|telefone|phone|celular|"
            r"endereco|address|objetivo|summary|perfil|github",
            re.IGNORECASE,
        )
        name = ""
        for line in clean_lines[:25]:
            if len(line) < 3 or len(line) > 90:
                continue
            if skip_name_patterns.search(line):
                continue
            if re.search(r"[@:/\\]|\d{3,}", line):
                continue
            words = [word for word in re.split(r"\s+", line) if word]
            if 2 <= len(words) <= 8:
                name = line
                break

        if not name and email_match:
            name = email_match.group(0).split("@", 1)[0].replace(".", " ").replace("_", " ").title()

        if not name and not email_match and not phone_match:
            self._set_error(
                "fallback_local_sem_identificador",
                f"Extracao local nao encontrou nome, e-mail ou telefone em {file_path}",
            )
            return None

        linkedin = self._normalizar_linkedin(linkedin_match.group(0) if linkedin_match else "")
        phone = phone_match.group(0).strip() if phone_match else ""
        email = email_match.group(0).strip().lower() if email_match else ""

        section_groups = {
            "perfil": [
                "perfil", "perfil profissional", "resumo", "resumo profissional",
                "summary", "professional summary", "objetivo", "objetivo profissional",
                "qualificacoes", "qualificacoes profissionais", "sobre mim"
            ],
            "formacao": [
                "formacao", "formacao academica", "educacao", "education",
                "escolaridade", "academic background", "graduacao"
            ],
            "cursos": [
                "cursos", "certificacoes", "certificados", "treinamentos",
                "cursos e certificacoes", "cursos certificacoes",
                "certificacoes e cursos", "treinamentos e certificacoes",
                "courses", "certifications", "certificates"
            ],
            "tecnico": [
                "conhecimentos tecnicos", "conhecimento tecnico", "competencias tecnicas",
                "habilidades tecnicas", "conhecimentos e habilidades tecnicas",
                "habilidades e conhecimentos tecnicos", "tecnologias", "ferramentas", "skills",
                "technical skills", "hard skills", "stack", "sistemas"
            ],
            "experiencia": [
                "experiencia", "experiencia profissional", "experiencias profissionais",
                "historico profissional", "trajetoria profissional", "professional experience",
                "work experience", "experiencia corporativa", "atuacao profissional"
            ],
            "idiomas": [
                "idiomas", "languages"
            ],
            "dados": [
                "dados pessoais", "informacoes pessoais", "contato", "contact"
            ],
        }
        all_headings = [alias for aliases in section_groups.values() for alias in aliases]

        perfil = self._extract_section(clean_lines, section_groups["perfil"], all_headings, max_chars=2500)
        formacao = self._extract_section(clean_lines, section_groups["formacao"], all_headings, max_chars=4000)
        cursos = self._extract_section(clean_lines, section_groups["cursos"], all_headings, max_chars=4000)
        tecnico = self._extract_section(clean_lines, section_groups["tecnico"], all_headings, max_chars=5000)
        experiencia = self._extract_section(clean_lines, section_groups["experiencia"], all_headings, max_chars=12000)

        tech_keywords = [
            "sap", "totvs", "protheus", "rm", "fluig", "datasul", "advpl", "sql",
            "java", "javascript", "typescript", "python", "c#", ".net", "react",
            "angular", "node", "azure", "aws", "power bi", "scrum", "kanban",
            "jira", "oracle", "mongodb", "docker", "kubernetes", "api", "erp"
        ]
        if not tecnico:
            tecnico = self._extract_keyword_lines(clean_lines, tech_keywords, max_lines=45, max_chars=5000)

        if not perfil:
            perfil = tecnico[:1600].strip() or compact_text[:1600].strip()

        if not experiencia:
            experiencia = compact_text[:12000].strip()

        estado_civil = self._extract_label_value(clean_lines, ["estado civil", "civil status"])
        nacionalidade = self._extract_label_value(clean_lines, ["nacionalidade", "nationality"])
        idade = self._extract_label_value(clean_lines, ["idade", "age"])
        endereco = self._extract_label_value(clean_lines, ["endereco", "endereço", "cidade", "localizacao", "localização", "address", "location"])

        if not idade:
            idade_match = re.search(r"\b(\d{2})\s+anos\b", compact_text, re.IGNORECASE)
            idade = idade_match.group(1) if idade_match else ""

        ingles = self._extract_language_level(clean_lines, ["ingles", "inglês", "english"])
        espanhol = self._extract_language_level(clean_lines, ["espanhol", "spanish"])

        return {
            "Nome": name,
            "Nacionalidade": nacionalidade,
            "Estado_Civil": estado_civil,
            "Idade": idade,
            "Endereco": endereco,
            "Telefone": phone,
            "Email": email,
            "Link_Linkedin": linkedin,
            "Skil": perfil,
            "Formacao_Academica": formacao,
            "Nivel_Idioma_Ingles": ingles,
            "Nivel_Idioma_Espanhol": espanhol,
            "Cursos_Certificacoes": cursos,
            "Conhecimento_Tecnico": tecnico,
            "Experiencia_Profissional": experiencia,
        }

    def extract_with_local_fallback(self, file_path):
        try:
            text = self._extract_text_local(file_path)
            dados = self._fallback_extract_from_text(text, file_path)
            if dados:
                logger.warning(
                    f"CV processado com extracao local de contingencia: {file_path}"
                )
                return dados
        except Exception as e:
            tipo_erro = self._classificar_exception(e)
            self._set_error(tipo_erro, f"Erro na extracao local de {file_path}: {e}")
            logger.error(f" Erro na extracao local do arquivo: {e}")

        return None
    
    def extract_from_file(self, file_path):
        """
        Extrai dados de um arquivo (PDF, DOC ou DOCX)
        
        Args:
            file_path (str): Caminho do arquivo
            
        Returns:
            dict: Dados extraídos ou None em caso de erro
        """
        
        try:
            # Garantir que seja PDF
            pdf_path = FileHandler.ensure_pdf(file_path)
            
            if not pdf_path:
                self._set_error(
                    "falha_conversao_pdf",
                    f"Não foi possível converter o arquivo para PDF: {file_path}"
                )
                logger.error(f" Não foi possível converter arquivo para PDF: {file_path}")
                return self.extract_with_local_fallback(file_path)
            
            # Extrair dados
            dados = self.extract_from_pdf(pdf_path)
            
            # Limpar arquivo temporário se foi convertido
            if pdf_path != file_path:
                FileHandler.delete_file(pdf_path)
            
            return dados or self.extract_with_local_fallback(file_path)
            
        except Exception as e:
            tipo_erro = self._classificar_exception(e)

            self._set_error(
                tipo_erro,
                f"Erro ao extrair dados do arquivo {file_path}: {e}"
            )

            logger.error(f" Erro ao extrair dados do arquivo: {e}")
            return self.extract_with_local_fallback(file_path)


def extract_cv_data_detalhado(file_path):
    """
    Extrai dados do CV retornando sucesso ou erro detalhado.
    """
    extractor = GeminiExtractor()
    dados = extractor.extract_from_file(file_path)

    if dados:
        return {
            "success": True,
            "dados": dados,
            "erro_tipo": None,
            "erro_mensagem": None
        }

    erro = extractor.last_error or {
        "tipo": "falha_extracao_sem_detalhe",
        "mensagem": "Não foi possível extrair dados do currículo."
    }

    return {
        "success": False,
        "dados": None,
        "erro_tipo": erro.get("tipo"),
        "erro_mensagem": erro.get("mensagem")
    }


def extract_cv_data(file_path):
    """
    Mantém compatibilidade com o código antigo.
    """
    resultado = extract_cv_data_detalhado(file_path)
    return resultado["dados"] if resultado["success"] else None
